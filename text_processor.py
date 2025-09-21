# optimized_text_processor.py
import re
from typing import List, Tuple, Dict, Set
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sentence_transformers import SentenceTransformer
import json

# Alternative PDF processing options
try:
    import fitz  # PyMuPDF
    PDF_LIBRARY = "pymupdf"
except ImportError:
    try:
        from PyPDF2 import PdfReader
        PDF_LIBRARY = "pypdf2"
    except ImportError:
        try:
            import pypdf
            PDF_LIBRARY = "pypdf"
        except ImportError:
            PDF_LIBRARY = None

from docx import Document
import spacy

class EnhancedTextProcessor:
    def __init__(self):
        # Load pre-trained models
        self.sentence_model = SentenceTransformer('all-MiniLM-L6-v2')
        
        # Try to load spaCy model
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            print("Warning: spaCy model not found. Using basic processing.")
            self.nlp = None
        
        # Enhanced skills database based on your sample documents
        self.technical_skills = {
            # Programming Languages
            'python', 'sql', 'r', 'java', 'javascript', 'c++', 'c#', 'scala', 'matlab',
            
            # Python Libraries (from your resume)
            'pandas', 'numpy', 'matplotlib', 'seaborn', 'scikit-learn', 'scipy',
            'beautiful soup', 'requests', 'regex', 'plotly', 'streamlit', 'flask', 'django',
            
            # Databases
            'mysql', 'postgresql', 'mongodb', 'sqlite', 'oracle', 'sql server',
            
            # Data Tools
            'power bi', 'tableau', 'excel', 'pivot tables', 'dax', 'power query',
            'jupyter', 'google analytics', 'looker', 'qlik',
            
            # Cloud & Big Data
            'aws', 'azure', 'gcp', 'spark', 'hadoop', 'kafka', 'airflow',
            'snowflake', 'redshift', 'bigquery',
            
            # AI/ML (from job description)
            'machine learning', 'deep learning', 'nlp', 'llm', 'generative ai',
            'tensorflow', 'pytorch', 'keras', 'transformers', 'bert', 'gpt',
            
            # Manufacturing/Engineering (from job description)
            'manufacturing', 'production', 'automotive', 'aerospace', 'mechanical engineering',
            'quality control', 'lean manufacturing', 'six sigma', 'process optimization',
            
            # Data Analysis
            'statistics', 'statistical analysis', 'hypothesis testing', 'correlation analysis',
            'descriptive statistics', 'data visualization', 'eda', 'exploratory data analysis',
            'data cleaning', 'data transformation', 'data mining', 'predictive modeling',
            
            # Business Intelligence
            'kpi', 'dashboard', 'reporting', 'business intelligence', 'data warehousing',
            'etl', 'data modeling', 'olap',
            
            # Soft Skills
            'problem solving', 'critical thinking', 'collaboration', 'communication',
            'project management', 'stakeholder management', 'agile', 'scrum'
        }
        
        # Industry-specific keywords for better matching
        self.domain_keywords = {
            'data_analysis': ['analysis', 'insights', 'trends', 'patterns', 'metrics', 'kpi'],
            'manufacturing': ['production', 'quality', 'process', 'automation', 'efficiency'],
            'ai_ml': ['artificial intelligence', 'machine learning', 'model', 'algorithm'],
            'business': ['stakeholder', 'business value', 'roi', 'optimization', 'strategy']
        }
    
    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """Enhanced PDF text extraction with multiple library support"""
        try:
            if PDF_LIBRARY == "pymupdf":
                return self._extract_with_pymupdf(file_content)
            elif PDF_LIBRARY == "pypdf2":
                return self._extract_with_pypdf2(file_content)
            elif PDF_LIBRARY == "pypdf":
                return self._extract_with_pypdf(file_content)
            else:
                raise Exception("No PDF library available. Please install PyMuPDF or PyPDF2.")
        except Exception as e:
            raise Exception(f"Error extracting PDF text: {str(e)}")
    
    def _extract_with_pymupdf(self, file_content: bytes) -> str:
        """Extract text using PyMuPDF (fitz)"""
        doc = fitz.open(stream=file_content, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text
    
    def _extract_with_pypdf2(self, file_content: bytes) -> str:
        """Extract text using PyPDF2"""
        from io import BytesIO
        pdf_reader = PdfReader(BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    
    def _extract_with_pypdf(self, file_content: bytes) -> str:
        """Extract text using pypdf"""
        from io import BytesIO
        import pypdf
        pdf_reader = pypdf.PdfReader(BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    
    def extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            from io import BytesIO
            doc = Document(BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
        except Exception as e:
            raise Exception(f"Error extracting DOCX text: {str(e)}")
    
    def clean_text(self, text: str) -> str:
        """Enhanced text cleaning and normalization"""
        # Remove extra whitespace and normalize
        text = re.sub(r'\s+', ' ', text)
        text = text.strip()
        
        # Handle common PDF extraction issues
        text = re.sub(r'[^\w\s\.\,\:\;\-\(\)\[\]\/\+\@]', ' ', text)
        
        # Normalize common variations
        text = re.sub(r'\bPython\b', 'python', text, flags=re.IGNORECASE)
        text = re.sub(r'\bSQL\b', 'sql', text, flags=re.IGNORECASE)
        text = re.sub(r'\bPower\s*BI\b', 'power bi', text, flags=re.IGNORECASE)
        text = re.sub(r'\bMachine\s*Learning\b', 'machine learning', text, flags=re.IGNORECASE)
        
        return text.lower()
    
    def extract_skills_enhanced(self, text: str) -> List[str]:
        """Enhanced skill extraction with better matching"""
        text_lower = text.lower()
        found_skills = []
        
        # Direct skill matching
        for skill in self.technical_skills:
            # Use word boundaries for better matching
            skill_pattern = r'\b' + re.escape(skill.lower()) + r'\b'
            if re.search(skill_pattern, text_lower):
                found_skills.append(skill.title())
        
        # Extract programming languages with special patterns
        prog_patterns = {
            'Python': r'\bpython\b|\bpandas\b|\bnumpy\b|\bmatplotlib\b',
            'SQL': r'\bsql\b|\bmysql\b|\bpostgresql\b|\bquery\b|\bqueries\b',
            'R': r'\br\b(?!\w)',  # R as standalone word
            'Excel': r'\bexcel\b|\bpivot\s*table\b|\bvlookup\b',
            'Power BI': r'\bpower\s*bi\b|\bdax\b|\bpower\s*query\b'
        }
        
        for skill, pattern in prog_patterns.items():
            if re.search(pattern, text_lower, re.IGNORECASE):
                found_skills.append(skill)
        
        # Remove duplicates while preserving order
        return list(dict.fromkeys(found_skills))
    
    def calculate_enhanced_hard_score(self, resume_text: str, job_text: str) -> Tuple[float, List[str]]:
        """Enhanced hard matching with better keyword analysis"""
        
        # Extract skills from both texts
        resume_skills = set(skill.lower() for skill in self.extract_skills_enhanced(resume_text))
        job_skills = set(skill.lower() for skill in self.extract_skills_enhanced(job_text))
        
        # Calculate skill-based score
        if not job_skills:
            skill_score = 50.0  # Default score if no skills detected
        else:
            overlap = len(resume_skills.intersection(job_skills))
            skill_score = (overlap / len(job_skills)) * 100
        
        # TF-IDF based matching for additional context
        try:
            # Create custom vocabulary focused on technical terms
            vectorizer = TfidfVectorizer(
                stop_words='english', 
                max_features=1000,
                ngram_range=(1, 2),  # Include bigrams
                min_df=1
            )
            
            tfidf_matrix = vectorizer.fit_transform([resume_text, job_text])
            cosine_sim = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
            tfidf_score = cosine_sim * 100
            
            # Weighted combination (60% skills, 40% TF-IDF)
            combined_score = (skill_score * 0.6) + (tfidf_score * 0.4)
        except:
            combined_score = skill_score
        
        # Find missing critical skills
        missing_skills = list(job_skills - resume_skills)
        
        # Prioritize missing skills based on frequency in job description
        prioritized_missing = []
        job_text_lower = job_text.lower()
        
        for skill in missing_skills:
            skill_count = len(re.findall(r'\b' + re.escape(skill) + r'\b', job_text_lower))
            prioritized_missing.append((skill, skill_count))
        
        # Sort by frequency and take top skills
        prioritized_missing.sort(key=lambda x: x[1], reverse=True)
        final_missing_skills = [skill.title() for skill, _ in prioritized_missing[:8]]
        
        return min(combined_score, 100.0), final_missing_skills
    
    def calculate_semantic_score(self, resume_text: str, job_text: str) -> float:
        """Enhanced semantic similarity calculation"""
        try:
            # Split longer texts into chunks for better processing
            resume_chunks = self._split_text(resume_text, max_length=500)
            job_chunks = self._split_text(job_text, max_length=500)
            
            # Calculate embeddings for chunks
            resume_embeddings = self.sentence_model.encode(resume_chunks)
            job_embeddings = self.sentence_model.encode(job_chunks)
            
            # Find best matching chunks
            similarities = []
            for job_emb in job_embeddings:
                chunk_similarities = cosine_similarity([job_emb], resume_embeddings)[0]
                similarities.append(max(chunk_similarities))
            
            # Average similarity across all job chunks
            avg_similarity = sum(similarities) / len(similarities)
            
            return min(avg_similarity * 100, 100.0)
            
        except Exception as e:
            print(f"Error calculating semantic similarity: {e}")
            # Fallback to simple word overlap
            return self._simple_word_overlap(resume_text, job_text)
    
    def _split_text(self, text: str, max_length: int = 500) -> List[str]:
        """Split text into meaningful chunks"""
        sentences = re.split(r'[.!?]+', text)
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            sentence = sentence.strip()
            if len(current_chunk) + len(sentence) < max_length:
                current_chunk += sentence + ". "
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence + ". "
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks if chunks else [text[:max_length]]
    
    def _simple_word_overlap(self, text1: str, text2: str) -> float:
        """Simple word overlap fallback"""
        words1 = set(re.findall(r'\b\w+\b', text1.lower()))
        words2 = set(re.findall(r'\b\w+\b', text2.lower()))
        
        if not words2:
            return 0.0
        
        overlap = len(words1.intersection(words2))
        return (overlap / len(words2)) * 100
    
    def generate_enhanced_feedback(self, hard_score: float, soft_score: float, 
                                 missing_skills: List[str], resume_text: str, job_text: str) -> str:
        """Generate comprehensive, personalized feedback"""
        
        feedback_parts = []
        
        # Overall assessment
        overall_score = (hard_score * 0.7) + (soft_score * 0.3)
        
        if overall_score >= 80:
            feedback_parts.append("🎯 Excellent match! Your profile strongly aligns with this role.")
        elif overall_score >= 65:
            feedback_parts.append("✅ Good match! You have most of the required qualifications.")
        elif overall_score >= 45:
            feedback_parts.append("⚠️ Moderate match. Consider highlighting relevant experience more prominently.")
        else:
            feedback_parts.append("🔄 This role requires significant skill development to be a strong match.")
        
        # Technical skills assessment
        if hard_score >= 70:
            feedback_parts.append("Your technical skills are well-aligned with the requirements.")
        elif hard_score >= 40:
            feedback_parts.append("You have some relevant technical skills, but there are gaps to address.")
        else:
            feedback_parts.append("Focus on developing the core technical skills for this role.")
        
        # Semantic alignment
        if soft_score >= 70:
            feedback_parts.append("Your experience narrative aligns well with the job context.")
        else:
            feedback_parts.append("Consider restructuring your resume to better reflect the job's language and requirements.")
        
        # Missing skills guidance
        if missing_skills:
            critical_skills = missing_skills[:3]  # Top 3 most important
            if len(critical_skills) == 1:
                feedback_parts.append(f"Priority skill to develop: {critical_skills[0]}.")
            else:
                feedback_parts.append(f"Priority skills to develop: {', '.join(critical_skills[:-1])}, and {critical_skills[-1]}.")
            
            if len(missing_skills) > 3:
                feedback_parts.append(f"Additional skills to consider: {', '.join(missing_skills[3:6])}.")
        
        # Industry-specific advice based on job content
        job_lower = job_text.lower()
        if 'manufacturing' in job_lower or 'production' in job_lower:
            feedback_parts.append("💡 For manufacturing roles, emphasize process improvement and data-driven optimization experience.")
        
        if 'machine learning' in job_lower or 'ai' in job_lower:
            feedback_parts.append("🤖 Highlight any AI/ML projects, even academic ones, and consider showcasing model deployment experience.")
        
        if 'stakeholder' in job_lower or 'business' in job_lower:
            feedback_parts.append("👥 Emphasize your ability to translate technical insights into business value and stakeholder communication.")
        
        # Action items
        if overall_score < 60:
            feedback_parts.append("📋 Recommended actions: Update your skills section, add relevant projects, and use keywords from the job description.")
        
        return " ".join(feedback_parts)
    
    # Keep the old method names for compatibility
    def calculate_hard_match_score(self, resume_text: str, job_text: str) -> Tuple[float, List[str]]:
        """Compatibility wrapper for enhanced hard matching"""
        return self.calculate_enhanced_hard_score(resume_text, job_text)
    
    def calculate_soft_match_score(self, resume_text: str, job_text: str) -> float:
        """Compatibility wrapper for semantic scoring"""
        return self.calculate_semantic_score(resume_text, job_text)
    
    def generate_feedback(self, hard_score: float, soft_score: float, missing_skills: List[str]) -> str:
        """Compatibility wrapper - simplified version"""
        return self.generate_enhanced_feedback(hard_score, soft_score, missing_skills, "", "")